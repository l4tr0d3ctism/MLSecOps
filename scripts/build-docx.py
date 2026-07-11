#!/usr/bin/env python3
"""Build MLSecOps Practical Reference Guide DOCX from chapters-en/*.md.

Uses the existing Word style template (reference DOCX) via Pandoc, embeds diagram
PNGs, and writes dist/MLSecOps-Practical-Reference-Guide-v{version}.docx.

Example (from repository root):

  pip install -r scripts/requirements-docx.txt
  python scripts/build-docx.py --render-mermaid

Diagram rendering uses Playwright with the system Chrome or Edge browser
(no separate `playwright install` required when Chrome/Edge is installed).
"""
from __future__ import annotations

import argparse
import os
import re
import shutil
import sys
import tempfile
import urllib.request
import zipfile
from pathlib import Path

try:
    import pypandoc
    from docx import Document
except ImportError:
    print("Install dependencies: pip install -r scripts/requirements-docx.txt", file=sys.stderr)
    sys.exit(1)

ROOT = Path(__file__).resolve().parent.parent
SCRIPTS = Path(__file__).resolve().parent
CHAPTERS = ROOT / "chapters-en"
DIAGRAMS = ROOT / "assets" / "diagrams"
DIAGRAM_SOURCE = DIAGRAMS / "source"
CITATION = ROOT / "CITATION.cff"
DIST = ROOT / "dist"
TEMPLATE_DIR = SCRIPTS / "templates"
DEFAULT_TEMPLATE = TEMPLATE_DIR / "reference.docx"
TEMPLATE_URL = (
    "https://github.com/l4tr0d3ctism/MLSecOps/releases/download/v1.0.0/"
    "MLSecOps-Practical-Reference-Guide-v1.0.0.docx"
)
IMAGE_WIDTH = "6.3in"

CHAPTER_FILES = [
    "01-intro.md",
    "02-scope-risk-threat-model.md",
    "03-threat-landscape.md",
    "04-data-security-privacy.md",
    "05-model-artifact-supply-chain.md",
    "06-pipeline.md",
    "07-llm-rag-security.md",
    "08-agentic-ai-security.md",
    "09-anti-patterns.md",
    "10-monitoring-soc-ir.md",
    "11-governance-evidence.md",
    "12-threat-control-tools-map.md",
    "13-case-studies.md",
    "14-maturity-roadmap.md",
    "15-conclusion-appendix.md",
    "16-kubernetes-deployment-reference.md",
    "17-appendix-e-implementation-reference.md",
]

PNG_REF = re.compile(r"!\[[^\]]*\]\(\.\./assets/diagrams/([^)]+)\)")
MERMAID_BLOCK = re.compile(r"```mermaid\s*\n(.*?)\n```", re.DOTALL)


def read_version() -> str:
    if not CITATION.exists():
        raise FileNotFoundError(f"Missing {CITATION}")
    for line in CITATION.read_text(encoding="utf-8").splitlines():
        if line.startswith("version:"):
            return line.split(":", 1)[1].strip().strip('"')
    raise RuntimeError("version not found in CITATION.cff")


def output_name(version: str) -> str:
    return f"MLSecOps-Practical-Reference-Guide-v{version}.docx"


def legacy_template_candidates() -> list[Path]:
    return [
        ROOT / "MLSecOps-Guide-v0.1.pre-v0.1.2-backup.docx",
        ROOT / "MLSecOps-Practical-Reference-Guide-v1.0.0.docx",
        ROOT / "MLSecOps-Guide-v0.1.docx",
    ]


def download_template(dest: Path) -> Path:
    dest.parent.mkdir(parents=True, exist_ok=True)
    print(f"Downloading reference template -> {dest}")
    urllib.request.urlretrieve(TEMPLATE_URL, dest)
    if not dest.exists() or dest.stat().st_size < 100_000:
        raise RuntimeError(f"Template download failed or file too small: {dest}")
    return dest


def resolve_reference(explicit: Path | None) -> Path:
    if explicit:
        if not explicit.exists():
            raise FileNotFoundError(f"Reference DOCX not found: {explicit}")
        return explicit.resolve()

    env = os.environ.get("MLSECOPS_REFERENCE_DOCX")
    if env:
        path = Path(env)
        if path.exists():
            return path.resolve()

    if DEFAULT_TEMPLATE.exists():
        return DEFAULT_TEMPLATE.resolve()

    for candidate in legacy_template_candidates():
        if candidate.exists():
            print(f"Using legacy reference template: {candidate.name}")
            return candidate.resolve()

    return download_template(DEFAULT_TEMPLATE)


def collect_required_pngs() -> set[str]:
    names: set[str] = set()
    for chapter in CHAPTER_FILES:
        text = (CHAPTERS / chapter).read_text(encoding="utf-8")
        names.update(PNG_REF.findall(text))
    return names


def ensure_diagram_pngs(render_mermaid: bool) -> None:
    required = collect_required_pngs()
    missing = [name for name in sorted(required) if not (DIAGRAMS / name).exists()]
    if not missing:
        print(f"Diagram PNGs OK ({len(required)} referenced)")
        return
    if not render_mermaid:
        raise FileNotFoundError(
            "Missing diagram PNGs in assets/diagrams/: "
            + ", ".join(missing[:5])
            + (" ..." if len(missing) > 5 else "")
            + ". Re-run with --render-mermaid or export PNGs from assets/diagrams/source/*.mmd"
        )

    if str(SCRIPTS) not in sys.path:
        sys.path.insert(0, str(SCRIPTS))
    from mermaid_to_png import mermaid_renderer_session, render_mermaid_to_png

    print(f"Rendering {len(missing)} missing diagram PNG(s) ...")
    with mermaid_renderer_session():
        for name in missing:
            src = DIAGRAM_SOURCE / name.replace(".png", ".mmd")
            if not src.exists():
                raise FileNotFoundError(f"No Mermaid source for {name}: expected {src}")
            out = DIAGRAMS / name
            render_mermaid_to_png(src.read_text(encoding="utf-8"), out)
            print(f"  {name}")


def md_for_docx(path: Path, assets_dir: Path) -> str:
    if str(SCRIPTS) not in sys.path:
        sys.path.insert(0, str(SCRIPTS))

    text = path.read_text(encoding="utf-8")

    if MERMAID_BLOCK.search(text):
        from mermaid_to_png import replace_mermaid_with_images

        text = replace_mermaid_with_images(text, assets_dir, path.stem)

    def png_repl(match: re.Match[str]) -> str:
        filename = match.group(1)
        png = (DIAGRAMS / filename).resolve()
        if not png.exists():
            raise FileNotFoundError(f"Diagram PNG not found: {png}")
        return f"\n\n![]({png.as_posix()}){{width={IMAGE_WIDTH}}}\n\n"

    text = PNG_REF.sub(png_repl, text)
    text = re.sub(r"(?<!!)\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\s*\{#[^}]+\}", "", text)
    return text


def build_full_markdown(version: str, assets_dir: Path) -> str:
    parts = [
        f"# MLSecOps Practical Reference Guide v{version}",
        "",
        "Practical reference guide — securing AI systems across the lifecycle. "
        f"Built from chapters-en/ (v{version}).",
        "",
        "\\newpage",
        "",
    ]
    for name in CHAPTER_FILES:
        parts.append(md_for_docx(CHAPTERS / name, assets_dir))
        parts.append("")
        parts.append("\\newpage")
        parts.append("")
    return "\n".join(parts)


def pandoc_build(md_text: str, reference: Path, assets_dir: Path, output: Path) -> None:
    tmp_dir = Path(tempfile.mkdtemp())
    try:
        md_file = tmp_dir / "MLSecOps-Guide.md"
        md_file.write_text(md_text, encoding="utf-8")
        resource_paths = [
            assets_dir.as_posix(),
            DIAGRAMS.as_posix(),
            ROOT.as_posix(),
            CHAPTERS.as_posix(),
        ]
        pypandoc.convert_file(
            str(md_file),
            "docx",
            format="md",
            outputfile=str(output),
            extra_args=[
                f"--reference-doc={reference}",
                f"--resource-path={os.pathsep.join(resource_paths)}",
            ],
        )
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def count_embedded_media(docx_path: Path) -> int:
    with zipfile.ZipFile(docx_path) as zf:
        return sum(
            1
            for name in zf.namelist()
            if name.startswith("word/media/") and name.lower().endswith(".png")
        )


def validate_docx(docx_path: Path, version: str, min_pngs: int) -> None:
    doc = Document(str(docx_path))
    full = "\n".join(p.text for p in doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full += "\n" + cell.text

    media_pngs = count_embedded_media(docx_path)
    print(f"Embedded PNG media files: {media_pngs}")
    if media_pngs < min_pngs:
        raise RuntimeError(f"Expected >= {min_pngs} embedded diagram PNGs, found {media_pngs}")

    must_present = [
        "References / Source mapping",
        "lifecycle control model",
        "Appendix E",
        "Chapter 16",
        f"v{version}",
    ]
    print("Validation:")
    for term in must_present:
        ok = term in full
        print(f"  present {term!r}: {'OK' if ok else 'MISSING'}")
        if not ok:
            raise RuntimeError(f"Missing expected content: {term}")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--version", help="Guide version (default: CITATION.cff)")
    parser.add_argument(
        "--output",
        type=Path,
        help="Output DOCX path (default: dist/MLSecOps-Practical-Reference-Guide-v{version}.docx)",
    )
    parser.add_argument(
        "--reference",
        type=Path,
        help="Word reference/template DOCX for Pandoc styles",
    )
    parser.add_argument(
        "--render-mermaid",
        action="store_true",
        help="Render missing diagram PNGs from assets/diagrams/source/*.mmd (requires Playwright)",
    )
    parser.add_argument(
        "--skip-validate",
        action="store_true",
        help="Skip post-build content checks",
    )
    parser.add_argument(
        "--min-pngs",
        type=int,
        default=20,
        help="Minimum embedded PNG count for validation (default: 20)",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    if not CHAPTERS.exists():
        print(f"Missing chapters dir: {CHAPTERS}", file=sys.stderr)
        sys.exit(1)

    version = args.version or read_version()
    output = args.output or (DIST / output_name(version))
    output.parent.mkdir(parents=True, exist_ok=True)
    reference = resolve_reference(args.reference)
    assets_dir = ROOT / "_docx_assets" / "mermaid"
    assets_dir.mkdir(parents=True, exist_ok=True)

    ensure_diagram_pngs(render_mermaid=args.render_mermaid)

    print(f"Building markdown (v{version}) ...")
    if args.render_mermaid and MERMAID_BLOCK.search(
        "\n".join((CHAPTERS / n).read_text(encoding="utf-8") for n in CHAPTER_FILES)
    ):
        from mermaid_to_png import mermaid_renderer_session

        with mermaid_renderer_session():
            md_text = build_full_markdown(version, assets_dir)
    else:
        md_text = build_full_markdown(version, assets_dir)

    image_refs = md_text.count("![")
    print(f"Image references in markdown: {image_refs}")

    tmp_out = output.with_suffix(".building.docx")
    print(f"Pandoc -> {tmp_out.name} (reference: {reference.name}) ...")
    pandoc_build(md_text, reference, assets_dir, tmp_out)

    if not args.skip_validate:
        validate_docx(tmp_out, version, args.min_pngs)

    shutil.copy2(tmp_out, output)
    tmp_out.unlink(missing_ok=True)

    doc = Document(str(output))
    print(f"Saved {output} ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")


if __name__ == "__main__":
    main()
